from docxtpl import DocxTemplate
import docx
from subprocess import Popen, PIPE


class FIles_Acts():
    def __init__(self, file_path):
        self.file_path = file_path

    def read_file(self) -> str:
        return open(self.file_path).read()

    def run_file(self, file_path):
        proc = Popen(["python3", file_path])
        stdout, stderr = proc.communicate()
        return stdout


class Make_Sample():
    def __init__(self, picture_path: str, codes: int):
        self.doc = docx.Document("python/ReportGen/templates/5_lab.docx")
        self.codes = codes
        self.picture_path = picture_path
        self.context = {"student": "", "group": "",
                        "teacher": "", "lab_numb": "", "variant": ""}

    def add_Task(self):
        self.doc.add_picture(self.picture_path)

    def write_paragraph(self):
        for i in range(self.codes):
            self.context[f"code{i}"] = ""
            self.context[f"result{i}"] = ""
            self.doc.add_paragraph(f"code{i}")
            self.doc.add_paragraph(f"result{i}")


class Make_Done_Report():
    def __init__(self, file_name: str, count_codes: int, programs_paths: list):
        self.sample = Make_Sample()
        self.file_name = file_name
        self.count_codes = count_codes
        self.docxt = DocxTemplate("smaple.docx")

    def add_codes(self):
        for i, value in enumerate(self.count_codes):
            context[f"code{i}"] = self.sample.read_file(value)
            context[f"result{i}"] = self.sample.run_file(value)

    def add_Task(self):
        pass


doc = DocxTemplate("smaple.docx")
context = {}
